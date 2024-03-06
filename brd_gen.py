# import os
import openai
import sqlite3
from docx import Document
import openai
import zipfile
from flask import Flask
# from dotenv import load_dotenv

# load_dotenv()
# openai.api_key_path = os.getenv("OPENAI_API_KEY")

openai.api_key = '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'


def save_section_datain_DB(project_name, user_business_idea):
    # Connect to the SQLite database
    print('savesection_function_called')
    conn = sqlite3.connect('blueprint.db')
    c = conn.cursor()
    print('dbconnected')

    sections = [
        "Project Overview",
        "Purpose and Scope of this Specification",
        "Product/Service Description",
        "Product Deliverables",
        "User Scenarios/Use Cases",
        "User Characteristics",
        "Roles and Actors",
        "Project Exclusions",
        "Business Requirements",
        "Process Steps",
        "Data Requirements",
        "Decision Points",
        "Use Case Description",
        "User Roles and Actors",
        "Business Rules",
        "User Interfaces",
        "Exceptions and Error Handling",
        "Performance Considerations",
        "Manageability/Maintainability",
        "Monitoring",
        "Maintenance",
        "System Interface/Integration",
        "Network and Hardware Interfaces",
        "Security",
        "Protection",
        "Authorization and Authentication",
        "Data Management",
        "Assumptions",
        "Constraints",
        "Dependencies",
        "Cost estimate",
        "Deleted or Deferred Requirement",
        "Feedback and Confirmation",
    ]

    # Add 'project_name' and 'user_business_idea' columns
    column_names = ['project_name', 'user_business_idea'] + [section.replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '').replace('.', '').replace('-', '_').lower() for section in sections]

    # Create a table to store BRD responses if it doesn't exist
    # Each section will be represented as a separate column in the table
    c.execute('''CREATE TABLE IF NOT EXISTS brd_responses
             ({})'''.format(','.join(column_names)))

    # Commit changes to the database
    conn.commit()
    
    # Get content for "Project Overview" separately
    main_prompt_overview = {
        "role": "user",
        "content": f"Given the business idea:\n{user_business_idea}. Please provide a detailed explanation for the Project Overview section in BRD\n\n"
    }

    # Continue with the OpenAI API call including the prompt
    response_overview = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": "You are a business analyst who generates detailed Business Requirements Document (BRD)s for a technology product. Provide a comprehensive and detailed explanation for the requested section of the BRD."},
            main_prompt_overview,
        ],
        temperature=1.0,
        top_p=0,
        frequency_penalty=0,
        presence_penalty=0.0,
    )

    # Remove asterisks from the response content
    content_overview = response_overview.choices[0].message['content'].replace('*', '')

    # Insert or update the content for "Project Overview" section in the database
    c.execute("INSERT INTO brd_responses (project_name, user_business_idea, project_overview) VALUES (?, ?, ?)", (project_name, user_business_idea, content_overview))

    # Commit changes to the database
    conn.commit()

    for section in sections[1:]:
        main_prompt = {
            "role": "user",
            "content": f"Given the business idea:\n{user_business_idea}. Please provide a detailed explanation for the {section} section in BRD\n\n"
        }

        # Continue with the OpenAI API call including the prompt
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=[
                {"role": "system", "content": "You are a business analyst who generates detailed Business Requirements Document (BRD)s for a technology product. Provide a comprehensive and detailed explanation for the requested section of the BRD."},
                main_prompt,
            ],
            temperature=1.0,
            top_p=0,
            frequency_penalty=0,
            presence_penalty=0.0,
        )

        # Remove asterisks from the response content
        content = response.choices[0].message['content'].replace('*', '')
        print(content)
        # Check if the project name already exists in the database
        c.execute("SELECT project_name FROM brd_responses WHERE project_name = ?", (project_name,))
        existing_record = c.fetchone()

        # If the project name doesn't exist, insert a new record
        if not existing_record:
            # Split the content string into individual values for each section
            column_values = [project_name, user_business_idea] + [content_section.strip() for content_section in content.split('Project Overview\n\n')]
            # Generate a list of column names and values to be inserted
            column_names = ['project_name', 'user_business_idea'] + [section.replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '').replace('.', '').replace('-', '_').lower() for section in sections[:len(column_values)-2]]
            # Insert the record into the database
            c.execute("INSERT INTO brd_responses ({}) VALUES ({})".format(
                ', '.join(column_names), ', '.join(['?'] * len(column_names))), column_values)
        else:
            # If the project name already exists, update the content for the section
            update_query = "UPDATE brd_responses SET {section} = ? WHERE project_name = ?".format(
                section=section.replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '').replace('.', '').replace('-', '_').lower())
            c.execute(update_query, (content, project_name))

    # Commit changes to the database and close the connection
    conn.commit()
    conn.close()

    print("Data saved to the database")
    brd_create(project_name)
    return print('done')        

def brd_create(project_name):

    conn = sqlite3.connect('blueprint.db')
    c = conn.cursor()

    # Query the database to retrieve the content of the selected record
    c.execute("SELECT * FROM brd_responses WHERE project_name = ?", (project_name,))
    record = c.fetchone()

    if record:
        # Create a new Word document
        doc = Document()
        doc.add_heading('Business Requirements Document', level=1)

        # Retrieve the column names (sections) from the database table
        c.execute("PRAGMA table_info(brd_responses)")
        columns_info = c.fetchall()
        sections = [column[1] for column in columns_info if column[1] not in ['project_name', 'user_business_idea']]

        # Add content from the selected record to the document
        for section, content in zip(sections, record[2:]):  # Skipping first two columns
            # Remove special characters
            content_cleaned = ''.join(char if ord(char) < 128 else ' ' for char in content)

            # Add section title and content to the document
            doc.add_heading(section.replace('_', ' ').title(), level=2)
            doc.add_paragraph(content_cleaned)

        # Close the database connection
        conn.close()

        # Save the document
        # doc_filename = f"brd_response_for_{project_name}.docx"
        doc_filename = f"Gen_BRD_Responses.docx"
        doc.save(doc_filename)

        print(f"Docx file for project '{project_name}' saved as {doc_filename}")
    else:
        print(f"No record found for project '{project_name}'")

    return print('file_generated_successfully')

